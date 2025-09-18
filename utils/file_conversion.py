import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageFont
from docx import Document
import pandas as pd
import io
import streamlit as st
import os
import tempfile

def convert_pdf_to_images(file_stream, dpi: int = 220):
    """Converts a PDF file stream into a list of PIL Image objects.
    Uses higher DPI for sharper rendering to improve visual diff sensitivity.
    """
    # Compute zoom factor from DPI (PDF default is 72 DPI)
    zoom = max(1.0, dpi / 72.0)
    mat = fitz.Matrix(zoom, zoom)

    # Ensure we have a fresh readable buffer for PyMuPDF
    pdf_bytes = file_stream.read() if hasattr(file_stream, "read") else file_stream
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    images = []
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        pix = page.get_pixmap(matrix=mat, alpha=False)  # No alpha; consistent RGB
        img_bytes = pix.tobytes("png")
        # Ensure PIL opens from its own buffer to avoid lazy file handle issues
        pil_img = Image.open(io.BytesIO(img_bytes)).convert('RGB')
        images.append(pil_img)
    doc.close()
    return images

def convert_docx_to_images(file_stream):
    """Converts a DOCX file to images by rendering text content."""
    document = Document(file_stream)
    images = []
    
    # Extract text and create images
    full_text = []
    for para in document.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    if full_text:
        # Split content into pages (approximate)
        content = "\n".join(full_text)
        page_images = convert_text_to_images(content)
        images.extend(page_images)
    
    return images

def convert_xlsx_to_images(file_stream):
    """Converts an XLSX file to images by rendering spreadsheet content."""
    df = pd.read_excel(file_stream)
    
    # Convert DataFrame to text representation
    text_content = df.to_string(index=True)
    
    # Split into multiple pages if content is too long
    lines = text_content.split('\n')
    max_lines_per_page = 50
    
    images = []
    for i in range(0, len(lines), max_lines_per_page):
        page_lines = lines[i:i + max_lines_per_page]
        page_content = '\n'.join(page_lines)
        page_images = convert_text_to_images(page_content)
        images.extend(page_images)
    
    return images

def convert_text_to_images(content, max_width=800, font_size=16):
    """Converts text content to a list of PIL Image objects with pagination."""
    try:
        # Try to load a system font
        font = ImageFont.truetype("arial.ttf", font_size)
    except:
        try:
            font = ImageFont.truetype("DejaVuSans.ttf", font_size)
        except:
            font = ImageFont.load_default()
    
    lines = content.split('\n')
    line_height = font_size + 8
    margin = 20
    max_height = 1000  # Maximum page height
    
    images = []
    current_page_lines = []
    current_height = margin
    
    for line in lines:
        # Calculate line width to handle wrapping
        line_width = len(line) * (font_size * 0.6)  # Approximate character width
        
        if line_width > max_width - 2 * margin:
            # Split long lines
            words = line.split()
            current_line = ""
            for word in words:
                test_line = current_line + " " + word if current_line else word
                test_width = len(test_line) * (font_size * 0.6)
                if test_width > max_width - 2 * margin:
                    if current_line:
                        current_page_lines.append(current_line)
                        current_height += line_height
                        current_line = word
                    else:
                        current_page_lines.append(word)
                        current_height += line_height
                else:
                    current_line = test_line
            if current_line:
                current_page_lines.append(current_line)
                current_height += line_height
        else:
            current_page_lines.append(line)
            current_height += line_height
        
        # Check if we need to start a new page
        if current_height > max_height - margin:
            # Create image for current page
            img = create_text_image(current_page_lines, max_width, current_height, font, margin)
            images.append(img)
            
            # Start new page
            current_page_lines = []
            current_height = margin
    
    # Create image for remaining lines
    if current_page_lines:
        img = create_text_image(current_page_lines, max_width, current_height, font, margin)
        images.append(img)
    
    return images

def create_text_image(lines, width, height, font, margin):
    """Creates a PIL Image with rendered text."""
    img = Image.new('RGB', (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    
    y = margin
    for line in lines:
        draw.text((margin, y), line, fill=(0, 0, 0), font=font)
        y += font.size + 8
    
    return img

def convert_image_file(file_stream, file_type):
    """Converts image files to PIL Image objects."""
    img = Image.open(file_stream)
    
    # Convert to RGB if necessary
    if img.mode != 'RGB':
        img = img.convert('RGB')
    
    return [img]

def convert_txt_to_images(file_stream):
    """Converts a TXT file stream into a list of PIL Image objects."""
    content = file_stream.read().decode('utf-8')
    return convert_text_to_images(content)

def get_file_type(file_name):
    """Returns the file extension."""
    return file_name.split('.')[-1].lower()

def convert_to_comparable_format(uploaded_file):
    """Converts an uploaded file to a list of page images and extracted text."""
    file_type = get_file_type(uploaded_file.name)
    
    page_images = []
    extracted_text = ""

    try:
        if file_type == 'pdf':
            # Read once, reuse for images and text extraction
            pdf_bytes = uploaded_file.read()
            page_images = convert_pdf_to_images(io.BytesIO(pdf_bytes), dpi=240)
            # Extract text from PDF (reuse same bytes)
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                extracted_text += page.get_text() + "\n"
            doc.close()
            uploaded_file.seek(0)
            
        elif file_type == 'docx':
            page_images = convert_docx_to_images(uploaded_file)
            # Extract text from DOCX
            document = Document(uploaded_file)
            for para in document.paragraphs:
                extracted_text += para.text + "\n"
            uploaded_file.seek(0)
            
        elif file_type == 'xlsx':
            page_images = convert_xlsx_to_images(uploaded_file)
            # Extract text from XLSX
            df = pd.read_excel(uploaded_file)
            extracted_text = df.to_string()
            uploaded_file.seek(0)
            
        elif file_type in ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'gif']:
            page_images = convert_image_file(uploaded_file, file_type)
            # Try OCR for text extraction
            try:
                import pytesseract
                img = Image.open(uploaded_file)
                extracted_text = pytesseract.image_to_string(img)
            except ImportError:
                extracted_text = "OCR not available. Install pytesseract for text extraction from images."
            except Exception as e:
                extracted_text = f"OCR failed: {str(e)}"
            uploaded_file.seek(0)
            
        elif file_type == 'txt':
            extracted_text = uploaded_file.read().decode('utf-8')
            page_images = convert_txt_to_images(io.BytesIO(extracted_text.encode('utf-8')))
            
        elif file_type == 'rtf':
            # Basic RTF text extraction (removes RTF formatting)
            content = uploaded_file.read().decode('utf-8')
            # Simple RTF text extraction - remove RTF codes
            import re
            extracted_text = re.sub(r'\\[a-z]+\d*\s?', '', content)
            extracted_text = re.sub(r'[{}]', '', extracted_text)
            page_images = convert_text_to_images(extracted_text)
            
        elif file_type == 'csv':
            df = pd.read_csv(uploaded_file)
            extracted_text = df.to_string()
            page_images = convert_text_to_images(extracted_text)
            uploaded_file.seek(0)
            
        else:
            st.error(f"Unsupported file type: {file_type}")
            return [], ""

    except Exception as e:
        st.error(f"Error processing {file_type} file: {str(e)}")
        return [], ""

    return page_images, extracted_text

def get_supported_file_types():
    """Returns a list of supported file types."""
    return {
        'pdf': 'PDF Documents',
        'docx': 'Word Documents',
        'xlsx': 'Excel Spreadsheets',
        'txt': 'Text Files',
        'rtf': 'Rich Text Format',
        'csv': 'CSV Files',
        'jpg': 'JPEG Images',
        'jpeg': 'JPEG Images',
        'png': 'PNG Images',
        'bmp': 'Bitmap Images',
        'tiff': 'TIFF Images',
        'gif': 'GIF Images'
    }